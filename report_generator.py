import datetime
import io
import math
import os.path
from datetime import timedelta
from os import path

import numpy as np
import matplotlib.pyplot as plt
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK


class FormatterMixin:
    """
    Класс-миксин для предоставления методов форматирования различных типов данных
    """

    @staticmethod
    def great_or_less_string(a: int | float | np.int64 | datetime.time,
                             b: int | float | np.int64 | datetime.time) -> str:
        """
        Метод для генерирования строки, информирующей о том больше или меньше параметр a параметра b
        :param a: левый операнд
        :param b: правый операнд
        :return: строку больше|меньше|значительно/незначительно больше|значительно/незначительно меньше
        """
        state_flags = ['больше', 'меньше']
        ind = bool(a < b)
        if isinstance(a, int | float | np.int64) and isinstance(b, int | float | np.int64):
            data = (a, b)
        elif isinstance(a, datetime.time) and isinstance(b, datetime.time):
            d1 = timedelta(hours=a.hour, minutes=a.minute, seconds=a.second)
            d2 = timedelta(hours=b.hour, minutes=b.minute, seconds=b.second)
            data = (d1, d2)
        else:
            return NotImplemented

        ratio = max(data) / min(data)
        if ratio >= 1.5:
            return 'значительно ' + state_flags[ind]
        elif ratio <= 1.02:
            return 'незначительно ' + state_flags[ind]
        else:
            return state_flags[ind]

    @staticmethod
    def great_or_less_range(a: int | float, b: int | float) -> str:
        """
        Метод проверяет во сколько раз больше (меньше) параметр a параметра b
        :param a: левый операнд
        :param b: правый операнд
        :return: str - в N раз больше
        """
        if isinstance(a, int | float) and isinstance(b, int | float):
            data = (a, b)
        elif isinstance(a, datetime.time) and isinstance(b, datetime.time):
            d1 = timedelta(hours=a.hour, minutes=a.minute, seconds=a.second)
            d2 = timedelta(hours=b.hour, minutes=b.minute, seconds=b.second)
            data = (d1, d2)
        else:
            return NotImplemented

        ratio = round(max(data) / min(data), 2)
        if a > b:
            return f'в {ratio} раз больше'
        elif a < b:
            return f'в {ratio} раз меньше'
        else:
            return f'равно'

    @staticmethod
    def get_outliers_rows(df: pd.DataFrame, label: str, outliers_rate: int | float = 1.5, is_campaigns=False):
        """
        Метод получает строки с выбросами в определенных столбцах (label) из DataFrame
        :param df: объект DataFrame (данные из csv-файла)
        :param label: метка столбца, по которому происходит поиск выбросов
        :return: объект DataFrame, содержащий строки с наличием выбросов в столбце label
        """
        df = df.copy(deep=True)
        if not is_campaigns:
            # удаляем строку лендинг из выборки для всех наборов меток кроме CAMPAIGN_LABELS
            df = df.drop([0])
        # переводим время в число для сравнения
        df.time = df.time.apply(lambda t: (t.hour * 60 + t.minute) * 60 + t.second)
        df = df[df[label] > 0]
        # квартили распределения
        # если количество элементов больше 2 ищем квартиль на основе медиан
        if len(df[label]) > 2:
            quantiles = df[label].quantile([0.25, 0.50, 0.75], interpolation='midpoint')
        # иначе используем метод по умолчанию (linear)
        else:
            quantiles = df[label].quantile([0.25, 0.50, 0.75])

        # межквартириальный размах
        IQR = quantiles.iloc[2] - quantiles.iloc[0]
        # оставляем строки с выбросами по полю label
        pos_outliers = df[df[label] >= IQR * outliers_rate]
        normal_distribution = df[df[label].between(-abs(IQR * outliers_rate), IQR * outliers_rate)]
        neg_outliers = df[df[label] <= -abs(IQR * outliers_rate)]
        # возвращаем число секунд в объект datetime.time
        pos_outliers = pos_outliers.assign(time=pos_outliers['time'].apply(
            lambda ts: datetime.time(hour=ts // 3600 // 60, minute=ts % 3600 // 60, second=ts % 60)))
        normal_distribution = normal_distribution.assign(time=normal_distribution['time'].apply(
            lambda ts: datetime.time(hour=ts // 3600 // 60, minute=ts % 3600 // 60, second=ts % 60)))
        neg_outliers = neg_outliers.assign(time=neg_outliers['time'].apply(
            lambda ts: datetime.time(hour=ts // 3600 // 60, minute=ts % 3600 // 60, second=ts % 60)))
        return pos_outliers, normal_distribution, neg_outliers

    @staticmethod
    def number_formatter(num) -> str:
        """
        Метод форматирует число разделяя разрядность пробелами
        :param num:
        :return:
        """
        return f"{num:,}".replace(",", " ")

    @staticmethod
    def time_to_str(time: datetime.time):
        """
        Метод форматирует объект datetime.time в удобочитаемую строку
        :param time:
        :return:
        """
        if time.hour:
            return time.strftime('%H:%M:%S')
        return time.strftime('%M:%S')

    @staticmethod
    def end_word_formatter(word: str, number: int) -> str:
        """
        Метод форматирует окончание слова в зависимости от числа
        :param word: маркер слова и объектов label
        :param number: число, в соотвествие с которым необходимо форматировать окончание
        :return:
        """
        forms = {
            'visits': ['визит', 'визита', 'визитов'],
            'views': ['посетитель', 'посетителя', 'посетителей']
        }

        if number % 100 in range(11, 19):
            return forms[word][2]
        elif number % 10 == 1:
            return forms[word][0]
        elif number % 10 in [2, 3, 4]:
            return forms[word][1]
        else:
            return forms[word][2]


class Data:
    """
    Класс для считывания и первичного форматирования данных из csv-файлов
    """

    def __init__(self, cur_rk_path, org_path, prev_rk_path, groups_path, campaign_path):
        self.RK_LABELS = ['action', 'views', 'conv_views', 'visits', 'conv_visits', 'aborted', 'perc_aborted', 'depth',
                          'time',
                          'new_users_with_abort', 'perc_new_users_with_abort', 'new_users', 'perc_new_users']
        self.CAMPAIGN_LABELS = ['action', 'views', 'visits', 'aborted', 'perc_aborted', 'depth', 'time',
                                'new_users_with_abort', 'perc_new_users_with_abort', 'new_users', 'perc_new_users']
        self.ORG_LABELS = ['serivce', 'views', 'visists', 'perc_aborted', 'depth', 'time', 'perc_new_users']
        self.cur_rk_df = self.read_rk_csv(cur_rk_path)
        self.org_df = self.read_org_csv(org_path)
        self.prev_rk_df = self.read_rk_csv(prev_rk_path)
        self.groups_df = self.read_campaign_csv(groups_path)
        self.campaigns_df = self.read_campaign_csv(campaign_path)

    def read_rk_csv(self, filename: path):
        """
        Чтение данных из csv формата RK_LABELS
        :param filename: путь к файлу
        :return: объект pandas.DataFrame
        """
        if not filename or not os.path.exists(filename):
            return pd.DataFrame()

        rk_df = pd.read_csv(filename)
        if len(rk_df.values[0]) != len(self.RK_LABELS):
            return pd.DataFrame()
        rk_df.columns = self.RK_LABELS

        rk_df.conv_views = pd.to_numeric(rk_df.conv_views)
        rk_df.conv_views = rk_df.conv_views.apply(self.percent_formatter)

        rk_df.conv_visits = pd.to_numeric(rk_df.conv_visits)
        rk_df.conv_visits = rk_df.conv_visits.apply(self.percent_formatter)

        rk_df.perc_aborted = rk_df.perc_aborted.apply(self.percent_formatter)
        rk_df.depth = rk_df.depth.apply(self.float_formatter)
        rk_df.time = rk_df.time.apply(self.str_to_time)
        rk_df.perc_new_users = rk_df.perc_new_users.apply(self.percent_formatter)
        rk_df.perc_new_users_with_abort = rk_df.perc_new_users_with_abort.apply(self.percent_formatter)

        return rk_df

    def read_org_csv(self, filename: path):
        """
        Чтение данных из csv формата ORG_LABELS
        :param filename: путь к файлу
        :return: объект pandas.DataFrame
        """
        org_df = pd.read_csv(filename)
        org_df.columns = self.ORG_LABELS
        org_df.perc_aborted = org_df.perc_aborted.apply(self.percent_formatter)
        org_df.depth = org_df.depth.apply(self.float_formatter)
        org_df.time = org_df.time.apply(self.str_to_time)
        org_df.perc_new_users = org_df.perc_new_users.apply(self.percent_formatter)

        return org_df

    def read_campaign_csv(self, filename: path):
        """
        Чтение данных из csv формата CAMPAIGN_LABELS
        :param filename: путь к файлу
        :return: объект pandas.DataFrame
        """
        if os.path.exists(filename):
            campaign_df = pd.read_csv(filename)
            campaign_df.columns = self.CAMPAIGN_LABELS

            campaign_df.perc_aborted = campaign_df.perc_aborted.apply(self.percent_formatter)
            campaign_df.depth = campaign_df.depth.apply(self.float_formatter)
            campaign_df.time = campaign_df.time.apply(self.str_to_time)
            campaign_df.perc_new_users = campaign_df.perc_new_users.apply(self.percent_formatter)
            campaign_df.perc_new_users_with_abort = campaign_df.perc_new_users_with_abort.apply(self.percent_formatter)

            return campaign_df

    @staticmethod
    def percent_formatter(num):
        """
        Метод форматирует дробные числа в проценты
        :param num:
        :return:
        """
        return round(float(num) * 100, 2)

    @staticmethod
    def str_to_time(t):
        """
        Метод форматирует строку, соответствующего формата в объект datetime.time
        :param t:
        :return:
        """
        return datetime.datetime.strptime(t, '%H:%M:%S').time()

    @staticmethod
    def float_formatter(num):
        """
        Метод форматирует float-числа до 2 знаков после запятой
        :param num:
        :return:
        """
        return round(float(num), 2)


class SectionWriter(FormatterMixin):
    """
    Класс для записи пунктов в формирующийся документ (отчёт)
    """

    def __init__(self, document, cur_rk_df, org_df, prev_rk_df, groups_path, campaigns_path):
        self.document = document

        data = Data(cur_rk_df, org_df, prev_rk_df, groups_path, campaigns_path)
        self.cur_rk_df = data.cur_rk_df
        self.prev_rk_df = data.prev_rk_df
        self.org_df = data.org_df
        self.groups_df = data.groups_df
        self.campaigns_df = data.campaigns_df

    def write_general_section(self):
        """
        Общие показатели
        :return:
        """

        # ВИЗИТЫ
        self.document.add_paragraph('Общие показатели:', style='List Number')
        p = self.document.add_paragraph(style='List Bullet')
        p.add_run(f'Было привлечено ')
        p.add_run(f'{self.number_formatter(self.cur_rk_df.views.iloc[0])} уникальных посетителей').bold = True
        p.add_run(f' и совершено ')
        p.add_run(f'{self.number_formatter(self.cur_rk_df.visits.iloc[0])} визитов.').bold = True

        # ДОЛЯ ОТКАЗОВ
        p = self.document.add_paragraph(style='List Bullet')
        p.add_run('Доля отказов').bold = True
        if not self.prev_rk_df.empty:
            p.add_run(
                f' составила {self.cur_rk_df.perc_aborted.iloc[0]} % что ' \
                f'{self.great_or_less_string(self.prev_rk_df.perc_aborted.iloc[0], self.org_df.perc_aborted.iloc[0])}, ' \
                f'чем в органическом трафике ({self.org_df.perc_aborted.iloc[0]} %) и ' \
                f'{self.great_or_less_string(self.cur_rk_df.perc_aborted.iloc[0], self.prev_rk_df.perc_aborted.iloc[0])}, ' \
                f'чем в предыдущей РК ({self.prev_rk_df.perc_aborted.iloc[0]} %).'
            )
        else:
            p.add_run(
                f' составила {self.cur_rk_df.perc_aborted.iloc[0]} % что ' \
                f'{self.great_or_less_string(self.cur_rk_df.perc_aborted.iloc[0], self.org_df.perc_aborted.iloc[0])}, ' \
                f'чем в органическом трафике ({self.org_df.perc_aborted.iloc[0]} %).'
            )

        # ГЛУБИНА ПРОСМОТРА
        p_depth = self.document.add_paragraph(style='List Bullet')

        cur_depth = self.cur_rk_df.depth.iloc[0]
        org_depth = self.org_df.depth.iloc[0]

        cur_org_compare_list = [cur_depth, org_depth]
        p_depth.add_run('Глубина просмотра').bold = True
        if not self.prev_rk_df.empty:
            prev_depth = self.prev_rk_df.depth.iloc[0]
            cur_prev_compare_list = [cur_depth, prev_depth]
            p_depth.add_run(
                f' составляет в среднем {self.cur_rk_df.depth.iloc[0]} стр., это ' \
                f'{self.great_or_less_string(cur_depth, org_depth)} ' \
                f'(в {round(max(cur_org_compare_list) / min(cur_org_compare_list), 2)} раз), ' \
                f'чем по органике ({org_depth} стр.), и ' \
                f'{self.great_or_less_string(cur_depth, prev_depth)} ' \
                f'(в {round(max(cur_prev_compare_list) / min(cur_prev_compare_list), 2)} раз) '
                f'чем в предыдущей РК ({prev_depth} стр.).')
        else:
            p_depth.add_run(
                f' составляет в среднем {self.cur_rk_df.depth.iloc[0]} стр., это ' \
                f'{self.great_or_less_string(cur_depth, org_depth)} ' \
                f'(в {round(max(cur_org_compare_list) / min(cur_org_compare_list), 2)} раз), ' \
                f'чем по органике ({org_depth} стр.).')

        # ВРЕМЯ ПРОСМОТРА
        p_time = self.document.add_paragraph(style='List Bullet')

        cur_time = self.cur_rk_df.time.iloc[0]
        org_time = self.org_df.time.iloc[0]

        p_time.add_run('Среднее время на сайте').bold = True
        if not self.prev_rk_df.empty:
            prev_time = self.prev_rk_df.time.iloc[0]
            p_time.add_run(f' — {self.time_to_str(cur_time)} что {self.great_or_less_string(cur_time, org_time)}' \
                           f', чем у органического трафика ({self.time_to_str(org_time)}) и ' \
                           f'{self.great_or_less_string(cur_time, prev_time)}, чем у трафика ' \
                           f'с предыдущей РК ({self.time_to_str(prev_time)})')
        else:
            p_time.add_run(f' — {self.time_to_str(cur_time)} что {self.great_or_less_string(cur_time, org_time)}' \
                           f', чем у органического трафика ({self.time_to_str(org_time)}).')

        cur_perc_new_users = self.cur_rk_df.perc_new_users.iloc[0]
        org_perc_new_users = self.org_df.perc_new_users.iloc[0]

        p_new_users = self.document.add_paragraph(style='List Bullet')
        p_new_users.add_run('Доля новых посетителей').bold = True
        if not self.prev_rk_df.empty:
            prev_perc_new_users = self.prev_rk_df.perc_new_users.iloc[0]
            p_new_users.add_run(f' составила {cur_perc_new_users} % (c учётом отказов),' \
                                f' что {self.great_or_less_string(cur_perc_new_users, org_perc_new_users)} результата по ' \
                                f'органике ({org_perc_new_users} %) и {self.great_or_less_string(cur_perc_new_users, prev_perc_new_users)},' \
                                f' чем по предыдущей РК ({prev_perc_new_users} % с учётом отказов)')
        else:
            p_new_users.add_run(f' составила {cur_perc_new_users} % (c учётом отказов),' \
                                f' что {self.great_or_less_string(cur_perc_new_users, org_perc_new_users)} результата по ' \
                                f'органике ({org_perc_new_users} %).')

    def write_page_views_section(self):
        """
        Данные посещаемости страниц
        :return:
        """
        p2 = self.document.add_paragraph('Посещение страниц:', style='List Number')
        p2.paragraph_format.line_spacing = 1.5
        p2.style.font.size = Pt(12)

        # ПОСЕЩАЕМОСТЬ
        # замена NaN-значений на 0
        # self.cur_rk_df.fillna(0)
        zeros_actions = self.cur_rk_df[self.cur_rk_df['views'] == 0]
        for i in range(1, len(self.cur_rk_df)):
            item = self.cur_rk_df.iloc[i].replace(np.nan, 0)
            if item.views != 0:
                p = self.document.add_paragraph(style='List Bullet')
                if 'посещен' in item.action.lower():
                    p.add_run('Действие ')
                    p.add_run(f'«{item.action}»').bold = True
                    p.add_run(
                        f' {item.views} {self.end_word_formatter("views", item.views)}; {item.conv_views} % конверсия посетителей из лендинга; '
                        f'{item.perc_aborted} % доля отказов '
                        f'(относительно визитов); {item.depth} стр. глубина просмотра (в среднем, '
                        f'без учёта отказников); {self.time_to_str(item.time)} время просмотра (в среднем, без учёта отказников); '
                        f'{item.perc_new_users} % доля новых пользователей (с учётом отказов).')
                else:
                    p.add_run('Действие ')
                    p.add_run(f'«{item.action}»').bold = True
                    p.add_run(
                        f' {item.views} {self.end_word_formatter("views", item.views)}; {item.conv_views} % конверсия посетителей; '
                        f'{item.perc_aborted} % доля отказов '
                        f'(относительно визитов); {item.depth} стр. глубина просмотра (в среднем, '
                        f'без учёта отказников); {self.time_to_str(item.time)} время просмотра (в среднем, без учёта отказников);')

        p = self.document.add_paragraph(style='List Bullet')
        p.add_run('По действиям ')
        p.add_run(f', '.join(f'«{zeros_item}»' for zeros_item in zeros_actions['action'].tolist()))
        # p.add_run(f'«{item.action}» ').bold = True
        p.add_run(' посещений не зафиксировано.')

    def write_funnel_graph_section(self):
        """
        Графики-воронки выполнения целевых действий
        :return:
        """
        p3 = self.document.add_paragraph()
        p3.add_run().add_break(WD_BREAK.PAGE)
        p3 = self.document.add_paragraph(f'Диаграммы выполнения целевых действий:', style='List Number')

        picture = self.document.add_paragraph()

        blocks_dict = dict()

        for action in self.cur_rk_df.action[1:]:
            block_name = action.split(': ')[0]
            if block_name not in blocks_dict:
                blocks_dict[block_name] = [action]
            else:
                blocks_dict[block_name].append(action)

        colors = ["#a9d18e", "#ffc000", "#ed7d31", "#5b9bd5", "#4472c4"]
        if not [1 for item in blocks_dict.values() if len(item) > 2]:
            picture.paragraph_format.left_indent = Inches(0.5)
            picture.add_run('Недостаточно данных для построения диаграмм.').italic = True
            return

        for action in blocks_dict:
            if len(blocks_dict[action]) >= 2:
                df = self.cur_rk_df[self.cur_rk_df['action'].str.contains(action)].sort_values(by=['views'],
                                                                                               ascending=False)
                y = [1, 3.6]
                x1 = [8, 12]
                x2 = [6, 2]

                fig = plt.figure(figsize=(12, 8))

                for i in range(len(df)):
                    # cmap = plt.get_cmap('summer')
                    plt.fill_betweenx(y=y, x1=x1, x2=x2, color=colors[i % 5])
                    y = [i + 3 for i in y]
                    x1 = [i + 4.3 for i in x1]
                    x2 = [i - 4.3 for i in x2]
                plt.xticks([], [])
                plt.yticks([i for i in range(2, len(df) * 3, 3)], df["action"].apply(lambda s: s.split(': ')[1])[::-1],
                           wrap=True, fontsize=18)

                for y, value in zip([i for i in range(2, len(df) * 3, 3)],
                                    df["views"].apply(self.number_formatter)[::-1]):
                    plt.text(7.3, y, value, fontsize=17, fontweight="bold", color="white", ha="center")

                # plt.ylabel("Stages")

                plt.title(f'Воронка трафика "{action}"', loc="center", fontsize=18, fontweight="bold", pad=30)
                plt.subplots_adjust(left=0.3)

                img = io.BytesIO()
                plt.savefig(img)
                img.seek(0)
                picture.add_run().add_picture(img, width=Cm(16.2), height=Cm(10.8))

    def write_items_by_outliers(self, min_items_num: int, df: pd.DataFrame, label: str, is_campaign: bool,
                                outlier_rate: float, write_best: bool = True):
        """
        Метод для формирования групп (пунктов) 'наибольшие/наименьшие' параметры для действий или кампаний
        :param min_items_num: минимальное количество элементов в каждой из групп
        :param df: DataFrame на основе которого отбираются данные
        :param label: метка столбца, на основе которого отбираются данные
        :param is_campaign: флаг, сигнализирующий на основе какой сущности происходит отбор
        :param outlier_rate: множитель, отвечающий за величину отклонения данных, которые будут считаться выбросом
        :param write_best: флаг, сигнализирующий о виде искомых групп (наибольшие или наименьшие)
        :return: None
        """
        pos_outliers, normal, neg_outliers = self.get_outliers_rows(df, label, is_campaigns=is_campaign,
                                                                    outliers_rate=outlier_rate)
        pos_outliers_perc_abort, normal_perc_abort, neg_outliers_perc_abort = self.get_outliers_rows(df,
                                                                                                     'perc_aborted',
                                                                                                     is_campaigns=is_campaign,
                                                                                                     outliers_rate=outlier_rate)
        pos_outliers_time, normal_time, neg_outliers_time = self.get_outliers_rows(df, 'time', is_campaigns=is_campaign,
                                                                                   outliers_rate=outlier_rate)
        if write_best:
            cur_outliers = pos_outliers
            normal = normal.sort_values(by=label, ascending=False)
        else:
            cur_outliers = neg_outliers
            normal = normal.sort_values(by=label, ascending=True)

        num_items = len(cur_outliers)
        for i in range(num_items):
            item = cur_outliers.iloc[i]
            p = self.document.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(1)
            p.add_run(
                f'«{item.action}» ({self.number_formatter(item[label])} {self.end_word_formatter(label, item[label])}).')  # выброс

            # проверка на высокие показатели отказов + время
            if item.action in pos_outliers_perc_abort.action.values:
                p.add_run(
                    f' Так же, стоит отметить, что данное действие имеет относительно большую долю отказов ({item.perc_aborted} %).')
                if item.action in pos_outliers_time.action.values:
                    p.add_run(f' Но, так же, относительно большее время просмотра ({item.time}).')
                elif item.action in pos_outliers_time.action.values:
                    p.add_run(f' И относительно малое время просмотра ({item.time}).')

            # проверка на низкие показатели отказов + время
            elif item.action in neg_outliers_perc_abort.action.values:
                p.add_run(
                    f' Так же, стоит отметить, что данное действие имеет относительно малую долю отказов ({item.perc_aborted} %).')
                if item.action in pos_outliers_time.action.values:
                    p.add_run(f' И относительно большое время просмотра ({item.time}).')
                elif item.action in pos_outliers_time.action.values:
                    p.add_run(f' Но относительно малое время просмотра ({item.time}).')

            # если отказыв в пределах нормы, ищем выбросы для данного действия по времени
            elif item.action in pos_outliers_time.action.values:
                p.add_run(
                    f' Так же, стоит отметить, что данное действие имеет относительно большое время просмотра ({item.time}).')
            elif item.action in neg_outliers_time.action.values:
                p.add_run(
                    f' Так же, стоит отметить, что данное действие имеет относительно малое время просмотра ({item.time}).')

        if num_items < min_items_num:
            for i in range(len(normal[:min_items_num - num_items])):
                item = normal.iloc[i]
                p = self.document.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(1)
                p.add_run(
                    f'«{item.action}» ({self.number_formatter(item[label])} {self.end_word_formatter(label, item[label])}).')

                # проверка на высокие показатели отказов + время
                if item.action in pos_outliers_perc_abort.action.values:
                    p.add_run(
                        f' Так же, стоит отметить, что данное действие имеет относительно большую долю отказов ({item.perc_aborted} %).')
                    if item.action in pos_outliers_time.action.values:
                        p.add_run(f' Но относительно большее время просмотра ({item.time}).')
                    elif item.action in pos_outliers_time.action.values:
                        p.add_run(f' И относительно малое время просмотра ({item.time}).')

                # проверка на низкие показатели отказов + время
                elif item.action in neg_outliers_perc_abort.action.values:
                    p.add_run(
                        f' Так же, стоит отметить, что данное действие имеет относительно малую долю отказов ({item.perc_aborted} %).')
                    if item.action in pos_outliers_time.action.values:
                        p.add_run(f' И большее, относительно других действий, время просмотра ({item.time}).')
                    elif item.action in pos_outliers_time.action.values:
                        p.add_run(f' Но малое, относительно других действий, время просмотра ({item.time}).')

                # если отказыв в пределах нормы, ищем выбросы для данного действия по времени
                elif item.action in pos_outliers_time.action.values:
                    p.add_run(
                        f' Так же, стоит отметить, что данное действие имеет большое время просмотра ({item.time}).')
                elif item.action in neg_outliers_time.action.values:
                    p.add_run(
                        f' Так же, стоит отметить, что данное действие имеет малое время просмотра ({item.time}).')

    def write_outliers_section(self, outlier_rate: float):
        """
        Данные с анализом выбросов (если есть) в видеть наибольших или наименьших значений
        :param outlier_rate: множитель, отвечающий за величину отклонения данных, которые будут считаться выбросом
        :return:
        """
        self.document.add_paragraph(f'Анализ выбросов по действиям:', style='List Number')
        self.document.add_paragraph(
            'Наибольшее число визитов включают действия:').paragraph_format.left_indent = Inches(0.5)

        # четверть значений от общего числа - кол-во значений в минимальных или максимальных данных
        min_items_count = math.ceil(len(self.cur_rk_df) * 0.25)

        # записываем наибольшие по визитам строки
        self.write_items_by_outliers(min_items_count, self.cur_rk_df, 'visits', is_campaign=False,
                                     outlier_rate=outlier_rate)

        self.document.add_paragraph(
            'Наименьшее число визитов включают действия:').paragraph_format.left_indent = Inches(0.5)
        self.write_items_by_outliers(min_items_count, self.cur_rk_df, 'visits', is_campaign=False, write_best=False,
                                     outlier_rate=outlier_rate)

    def write_groups_section(self, outlier_rate: float):
        """
        данные с анализом групп, кампаний. Анализ схожий с пунктом write_outliers_section
        :param outlier_rate:
        :return:
        """
        self.document.add_paragraph(f'Группы:', style='List Number')
        if len(self.groups_df) == 2:
            p = self.document.add_paragraph(style='List Bullet')
            group_item1 = self.groups_df.iloc[0]
            group_item2 = self.groups_df.iloc[1]
            p.add_run(
                f'«{group_item1.action}» привлекла {self.great_or_less_string(group_item1.views, group_item2.views)} посетителей '
                f'({self.number_formatter(group_item1.views)}) и имеет '
                f'{self.great_or_less_string(group_item1.perc_aborted, group_item2.perc_aborted)} отказников '
                f'({group_item1.perc_aborted} %) чем группа «{group_item2.action}» ({self.number_formatter(group_item2.views)} '
                f'{self.end_word_formatter("views", group_item2.views)}; '
                f'{group_item2.perc_aborted} % отказников), и при этом имеет'
                f' {self.great_or_less_string(group_item1.time, group_item2.time)} среднее время на сайте '
                f'({self.time_to_str(group_item1.time)} у «{group_item1.action}» '
                f'против {self.time_to_str(group_item2.time)} у «{group_item2.action}»).'
            )
        elif len(self.groups_df) == 1:
            p = self.document.add_paragraph(style='List Bullet')
            group_item1 = self.groups_df.iloc[0]
            p.add_run(
                f'«{group_item1.action}» привлекла {self.number_formatter(group_item1.views)} '
                f'{self.end_word_formatter("views", group_item1.views)} '
                f'({self.number_formatter(group_item1.views)}) и имеет '
                f'{group_item1.perc_aborted} % отказников {self.time_to_str(group_item1.time)}'
            )
        else:
            min_items_count = math.ceil(len(self.groups_df) * 0.25)

            self.document.add_paragraph('Лучшие показатели посещаемости демонстируют следующие группы: ',
                                        style='List Bullet')
            self.write_items_by_outliers(min_items_count, self.groups_df, 'views', is_campaign=False, write_best=True,
                                         outlier_rate=outlier_rate)

            self.document.add_paragraph('Худшие показатели посещаемости демонстируют следующие группы: ',
                                        style='List Bullet')
            self.write_items_by_outliers(min_items_count, self.groups_df, 'views', is_campaign=False, write_best=False,
                                         outlier_rate=outlier_rate)

        min_items_count = math.ceil(len(self.campaigns_df) * 0.25)

        self.document.add_paragraph('Лучшие показатели посещаемости демонстрируют кампании:', style='List Bullet')
        self.write_items_by_outliers(min_items_count, self.campaigns_df, 'views', is_campaign=True,
                                     outlier_rate=outlier_rate)

        self.document.add_paragraph('Худшие показатели посещаемости у следующих кампаний:', style='List Bullet')
        self.write_items_by_outliers(min_items_count, self.campaigns_df, 'views', is_campaign=True, write_best=False,
                                     outlier_rate=outlier_rate)


class ReportGenerator(FormatterMixin):
    """
    Класс для управления формированием файла, создаёт объект документа, принимает пути к файлам с данными
    управляет записью пунктов в документ
    """

    def __init__(self, header: str, cur_rk_path: str, org_path: str, groups_path: str, campaigns_path: str,
                 prev_rk_path: str = None, outlier_rate: float = 1.5):
        """

        :param header: заголовок документа
        :param cur_rk_path: путь к файлу Текущая РК
        :param org_path: путь к файлу Органический трафик
        :param groups_path: путь к файлу с данными о группах
        :param campaigns_path: путь к файлу с данными о кампаниях
        :param prev_rk_path: путь к файлу с данными предыдущей РК
        :param outlier_rate: множитель, отвечающий за величину отклонения данных, которые будут считаться выбросом
        """
        self.document = Document()
        self.general_writer = SectionWriter(self.document, cur_rk_path, org_path, prev_rk_path, groups_path,
                                            campaigns_path)

        self.outlier_rate = outlier_rate

        self.write_header(header)
        # self.write_general_params()

        # настройки форматирования для заголовка
        self.header_style = self.document.styles['Normal']
        self.header_style.paragraph_format.line_spacing = 1.5
        header_font = self.header_style.font
        header_font.name = 'times new roman'
        header_font.size = Pt(12)

        # настройка форматирования для маркированного списка
        self.list_bullet_style = self.document.styles['List Bullet']
        self.list_bullet_style.paragraph_format.left_indent = Inches(0.5)
        self.list_bullet_style.paragraph_format.line_spacing = 1.5
        list_bullet_font = self.list_bullet_style.font
        list_bullet_font.name = 'times new roman'
        list_bullet_font.size = Pt(12)
        list_bullet_font.bold = False

    def write_header(self, header):
        header_obj = self.document.add_paragraph(style='Normal')
        header_obj.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_obj.add_run(f'Выводы по контекстной кампании "{header}"')
        run.font.size = Pt(14)
        run.bold = True
        header_obj.add_run('\n')

    def write_general_params(self):
        """
        Общие показатели
        :return:
        """
        self.general_writer.write_general_section()

    def write_page_views(self):
        """
        Просмотр страниц
        :return:
        """
        self.general_writer.write_page_views_section()

    def write_funnel_graph_section(self):
        """
        Графики-воронки
        :return:
        """
        self.general_writer.write_funnel_graph_section()

    def write_outliers_section(self):
        """
        Анализ выбросов, наилучших и наихудших параметров для действий
        :return:
        """
        self.general_writer.write_outliers_section(self.outlier_rate)

    def write_groups_section(self):
        """
        Анализ выбросов, наилучших и наихудших параметров для групп/кампаний
        :return:
        """
        self.general_writer.write_groups_section(self.outlier_rate)

    def save_report(self, doc_name: str = 'auto_report'):
        """
        Сохранение файла отчёта
        :param doc_name: имя выходного файла
        :return: None
        """
        self.document.save(doc_name)


if __name__ == '__main__':
    # объект для управления записью отчёта
    # report = ReportGenerator('Моя кампания', 'data/teatri_vov_um/Текущая РК.csv',
    #                          'data/teatri_vov_um/Органический трафик.csv',
    #                          'data/teatri_vov_um/Группы по типу РК.csv', 'data/teatri_vov_um/Все кампании.csv',
    #                          prev_rk_path='data/teatri_vov_um/Предыдущая РК.csv', outlier_rate=1.5)
    report = ReportGenerator('Моя кампания', 'data/old_data/Текущая РК.csv', 'data/old_data/Органический трафик.csv',
                             'data/old_data/Группы по типу РК.csv', 'data/old_data/Все кампании.csv',
                             prev_rk_path='data/old_data/Предыдущая РК.csv', outlier_rate=5)

    # вызов методов, для записи пунктов отчёта
    # общие показатели
    report.write_general_params()
    # посещение страниц
    report.write_page_views()
    # диаграммы выполнения целевых действий (воронки)
    report.write_funnel_graph_section()
    # анализ выбросов по действиям
    report.write_outliers_section()
    # анализ групп
    report.write_groups_section()

    # сохранение файла
    report.save_report('auto_report.docx')
