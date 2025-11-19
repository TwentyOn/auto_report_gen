import csv
import datetime
from datetime import timedelta
from os import path

import numpy
import numpy as np
import matplotlib.pyplot as plt
import random
from collections import Counter

import pandas as pd

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

RK_LABELS = ['action', 'views', 'conv_views', 'visits', 'conv_visits', 'aborted', 'perc_aborted', 'depth', 'time',
             'new_users_with_abort', 'perc_new_users_with_abort', 'new_users', 'perc_new_users']

ORG_LABELS = ['serivce', 'views', 'visists', 'perc_aborted', 'depth', 'time', 'perc_new_users']

# множитель нормального распределения, по которому из выборки выделяются выбросы, если = 2, то выбросы
# будут расчитываться от межквартилиального растояния * 2
OUTLIERS_RATE = 2


def great_or_less_string(a: int | float | datetime.time, b: int | float | datetime.time):
    state_flags = ['больше', 'меньше']
    ind = bool(a < b)
    if isinstance(a, int | float) and isinstance(b, int | float):
        data = (a, b)
    elif isinstance(a, datetime.time) and isinstance(b, datetime.time):
        d1 = timedelta(hours=a.hour, minutes=a.minute, seconds=a.second)
        d2 = timedelta(hours=b.hour, minutes=b.minute, seconds=b.second)
        data = (d1, d2)
    else:
        return NotImplemented

    ratio = max(data) / min(data)
    if ratio >= 1.5:
        return 'значительно ' + state_flags[ind]
    elif ratio <= 1.1:
        return 'незначительно' + state_flags[ind]
    else:
        return state_flags[ind]


def great_or_less_range(a: int | float, b: int | float):
    if isinstance(a, int | float) and isinstance(b, int | float):
        data = (a, b)
    elif isinstance(a, datetime.time) and isinstance(b, datetime.time):
        d1 = timedelta(hours=a.hour, minutes=a.minute, seconds=a.second)
        d2 = timedelta(hours=b.hour, minutes=b.minute, seconds=b.second)
        data = (d1, d2)
    else:
        print(type(a), a, type(b), b)
        return NotImplemented

    ratio = round(max(data) / min(data), 2)
    if a > b:
        return f'в {ratio} раз больше'
    elif a < b:
        return f'в {ratio} раз меньше'
    else:
        return f'равно'


def float_formatter(num):
    return round(float(num), 2)


def percent_formatter(num):
    return round(float(num) * 100, 2)


def str_to_time(t):
    return datetime.datetime.strptime(t, '%H:%M:%S').time()


def time_to_str(time: datetime.time):
    if time.hour:
        return time.strftime('%H:%M:%S')
    return time.strftime('%M:%S')


def get_outliers_rows(df: pd.DataFrame, label: str):
    """
    Функция получает строки с выбросами в определенных столбцах DataFrame помеченных label
    :param df: объект DataFrame (данные из csv-файла)
    :param label: метка столбца, по которому происходит поиск выбросов
    :return: объект DataFrame, содержащий строки с наличием выбросов в столбце label
    """
    # удаляем лендинг из выборки
    df = df.drop([0])
    # переводим время в число для сравнения
    df.time = df.time.apply(lambda t: (t.hour * 60 + t.minute) * 60 + t.second)
    # квартили распределения. Нулевые значения не учитываем
    quantiles = df[label].quantile([0.25, 0.50, 0.75])
    median = int(quantiles.iloc[1])
    if label == 'time':
        median = datetime.time(hour=median // 3600 // 60, minute=median % 3600 // 60, second=median % 60)
    # межквартириальный размах
    IQR = quantiles.iloc[2] - quantiles.iloc[0]
    # оставляем строки с выбросами по полю label, выбросом считается значение выше или ниже чем два IQR
    df = df[abs(df[label]) >= IQR * OUTLIERS_RATE]
    # приводим число в объект datetime.time
    df.time = df.time.apply(lambda ts: datetime.time(hour=ts // 3600 // 60, minute=ts % 3600 // 60, second=ts % 60))
    return df, median


def outliers_formatter(outliers_list: list[str], param):
    """
    Функция возвращает форматированную строку ответа для списка с выбросами
    :param outliers_values: список строк действие-параметр
    :param param_name: имя параметра
    :return: форматированная строка ответа
    """
    params = {
        'visists': 'Количество визитов',
        'conv_visist': 'Конверсия визитов',
        'perc_aborted': 'Доля отказов',
        'time': 'Время на сайте'
    }
    ans = f'По параметру «{param}» значительные расхождения данных замечены в следующих действиях: '

    return ans + '; '.join(outliers_list)


def read_rk_csv(filename: path):
    rk_df = pd.read_csv(filename)
    rk_df.columns = RK_LABELS

    rk_df.conv_views = pd.to_numeric(rk_df.conv_views)
    rk_df.conv_views = rk_df.conv_views.apply(percent_formatter)

    rk_df.conv_visits = pd.to_numeric(rk_df.conv_visits)
    rk_df.conv_visits = rk_df.conv_visits.apply(percent_formatter)

    rk_df.perc_aborted = rk_df.perc_aborted.apply(percent_formatter)
    rk_df.depth = rk_df.depth.apply(float_formatter)
    rk_df.time = rk_df.time.apply(str_to_time)
    rk_df.perc_new_users = rk_df.perc_new_users.apply(percent_formatter)
    rk_df.perc_new_users_with_abort = rk_df.perc_new_users_with_abort.apply(percent_formatter)

    return rk_df


def read_org_csv(filename: path):
    org_df = pd.read_csv(filename)
    org_df.columns = ORG_LABELS
    org_df.perc_aborted = org_df.perc_aborted.apply(percent_formatter)
    org_df.depth = org_df.depth.apply(float_formatter)
    org_df.time = org_df.time.apply(str_to_time)
    org_df.perc_new_users = org_df.perc_new_users.apply(percent_formatter)

    return org_df


def time_to_seconds(time: datetime.time):
    seconds = 0
    seconds += time.hour * 3600
    seconds += time.minute * 60
    seconds += time.second
    return seconds


def seconds_to_time(seconds):
    seconds = seconds.astype(np.int64)
    hour = seconds // 3600
    minute = seconds % 3600 // 60
    second = seconds % 60
    return datetime.time(hour=hour, minute=minute, second=second)


cur_rk_data = read_rk_csv('Текущая РК.csv')
prev_rk_data = read_rk_csv('Предыдущая РК.csv')
org_data = read_org_csv('Органический трафик.csv')

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'times new roman'
font.size = Pt(12)

# ЗАГОЛОВОК
header = document.add_paragraph(style=style)
header.add_run('Выводы по контекстной кампании "ИМЯ КАМПАНИИ"').bold = True
header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
header.style.font.name = 'times new roman'
header.style.font.size = Pt(14)
header.add_run('\n')

# ПУНКТ 1 ОБЩИЕ ПОКАЗАТЕЛИ
p1 = document.add_paragraph('Общие показатели:', style='List Number')
p1.style.font.size = Pt(12)
p1.paragraph_format.line_spacing = 1.5

# ВИЗИТЫ
p1_1 = document.add_paragraph(style='List Bullet')
p1_1.paragraph_format.line_spacing = 1.5
p1_1.style.font.size = Pt(12)
p1_1.paragraph_format.left_indent = Inches(0.5)
p1_1.add_run(f'Было привлечено ')
p1_1.add_run(f'{cur_rk_data.views.iloc[0]} уникальных посетителей').bold = True
p1_1.add_run(f' и совершено ')
p1_1.add_run(f'{cur_rk_data.visits.iloc[0]} визитов.').bold = True

# ДОЛЯ ОТКАЗОВ
p1_2 = document.add_paragraph(style='List Bullet')
p1_2.paragraph_format.line_spacing = 1.5
p1_2.style.font.size = Pt(12)
p1_2.paragraph_format.left_indent = Inches(0.5)
p1_2.add_run('Доля отказов').bold = True
p1_2.add_run(
    f' составила {cur_rk_data.perc_aborted.iloc[0]} % что ' \
    f'{great_or_less_string(cur_rk_data.perc_aborted.iloc[0], org_data.perc_aborted.iloc[0])}, ' \
    f'чем в органическом трафике ({org_data.perc_aborted.iloc[0]} %) и ' \
    f'{great_or_less_string(cur_rk_data.perc_aborted.iloc[0], prev_rk_data.perc_aborted.iloc[0])}, ' \
    f'чем в предыдущей РК ({prev_rk_data.perc_aborted.iloc[0]} %).'
)

# ГЛУБИНА ПРОСМОТРА
p1_3 = document.add_paragraph(style='List Bullet')
p1_3.paragraph_format.line_spacing = 1.5
p1_3.style.font.size = Pt(12)
p1_3.paragraph_format.left_indent = Inches(0.5)

cur_depth = cur_rk_data.depth.iloc[0]
org_depth = org_data.depth.iloc[0]
prev_depth = prev_rk_data.depth.iloc[0]

cur_org_compare_list = [cur_depth, org_depth]
cur_prev_compare_list = [cur_depth, prev_depth]
p1_3.add_run('Глубина просмотра').bold = True
p1_3.add_run(
    f' составляет в среднем {cur_rk_data.depth.iloc[0]} стр., это ' \
    f'{great_or_less_string(cur_depth, org_depth)} ' \
    f'(в {round(max(cur_org_compare_list) / min(cur_org_compare_list), 2)} раз), ' \
    f'чем по органике ({org_depth} стр.), и ' \
    f'{great_or_less_string(cur_depth, prev_depth)} ' \
    f'(в {round(max(cur_prev_compare_list) / min(cur_prev_compare_list), 2)} раз) '
    f'чем в предыдущей РК ({prev_depth} стр.).')

# ВРЕМЯ ПРОСМОТРА
p1_4 = document.add_paragraph(style='List Bullet')
p1_4.paragraph_format.line_spacing = 1.5
p1_4.style.font.size = Pt(12)
p1_4.paragraph_format.left_indent = Inches(0.5)

cur_time = cur_rk_data.time.iloc[0]
prev_time = prev_rk_data.time.iloc[0]
org_time = org_data.time.iloc[0]

p1_4.add_run('Среднее время на сайте').bold = True
p1_4.add_run(f' — {time_to_str(cur_time)} что {great_or_less_string(cur_time, prev_time)}' \
             f', чем у органического трафика ({time_to_str(org_time)}) и ' \
             f'{great_or_less_string(cur_time, org_time)}, чем у трафика ' \
             f'с предыдущей РК ({time_to_str(prev_time)})')

cur_perc_new_users = cur_rk_data.perc_new_users.iloc[0]
prev_perc_new_users = prev_rk_data.perc_new_users.iloc[0]
org_perc_new_users = org_data.perc_new_users.iloc[0]

p1_5 = document.add_paragraph(style='List Bullet')
p1_5.paragraph_format.line_spacing = 1.5
p1_5.style.font.size = Pt(12)
p1_5.paragraph_format.left_indent = Inches(0.5)
p1_5.add_run('Доля новых посетителей').bold = True
p1_5.add_run(f' составила {cur_perc_new_users} % (c учётом отказов),' \
             f' что {great_or_less_string(cur_perc_new_users, org_perc_new_users)} результата по ' \
             f'органике ({org_perc_new_users} %) и {great_or_less_string(cur_perc_new_users, prev_perc_new_users)},' \
             f' чем по предыдущей РК ({prev_perc_new_users} % с учётом отказов)')

# ПУНКТ 2 - Посещение страниц
p2 = document.add_paragraph('Посещение страниц:', style='List Number')
p2.paragraph_format.line_spacing = 1.5
p2.style.font.size = Pt(12)

# ПОСЕЩАЕМОСТЬ
cur_rk_data.fillna(0)

for i in range(len(cur_rk_data)):
    item = cur_rk_data.iloc[i].replace(np.nan, 0)
    p = document.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.5
    p.style.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(0.5)
    p.add_run(f'Действие «{item.action}»: {item.views} посещений; {item.conv_views} % конверсия посетителей; '
              f'{item.visits} визитов; {item.conv_visits} % конверсия визитов; '
              f'{item.aborted} количество отказов; {item.perc_aborted} % доля отказов '
              f'(относительно визитов); {item.depth} стр. глубина просмотра (в среднем, '
              f'без учёта отказников); {item.time} время просмотра (в среднем, без учёта отказников); '
              f'{item.new_users_with_abort} количество новых посетителей (без учёта отказов); '
              f'{item.perc_new_users_with_abort} % доля новых посетителей (без учёта отказов); '
              f'{item.new_users} колчество новых поситетелей (с учётом отказов); '
              f'{item.perc_new_users} % доля новых пользователей (с учётом отказов)')

# -----------------------------------------------------------------------------------------
# ПУНКТ 3 - ВЫБРОСЫ
outliers_visits, median_visits = get_outliers_rows(cur_rk_data, 'visits')
outliers_cov_visits, media_conv_visits = get_outliers_rows(cur_rk_data, 'conv_visits')
outliers_perc_abort, median_perc_abort = get_outliers_rows(cur_rk_data, 'perc_aborted')
outliers_time, median_time = get_outliers_rows(cur_rk_data, 'time')

p3 = document.add_paragraph(f'Выбросы данных (множитель нормального распределения = {OUTLIERS_RATE}):',
                              style='List Number')
p3.paragraph_format.line_spacing = 1.5
p3.style.font.size = Pt(12)

visits_outliers_list = [
    f"«{i[0]}» ({i[RK_LABELS.index('visits')]}), что {great_or_less_range(i[RK_LABELS.index('visits')], cur_rk_data.visits.mean())} чем в среднем ({cur_rk_data.visits.mean().round(2)}) при этом медиана для данного параметра = {median_visits}."
    for i in outliers_visits.values]
p3_1_1 = document.add_paragraph(style='List Bullet')
p3_1_1.paragraph_format.left_indent = Inches(0.5)
p3_1_1.paragraph_format.line_spacing = 1.5
p3_1_1.add_run(outliers_formatter(visits_outliers_list, 'Количество визитов'))
conv_visits_outliers_list = [
    f"«{i[0]}» ({i[RK_LABELS.index('conv_visits')]} %), что в {great_or_less_range(i[RK_LABELS.index('conv_visits')], cur_rk_data.conv_visits.mean())} чем в среднем ({cur_rk_data.conv_visits.mean().round(2)} %) при этом медиана для данного параметра = {media_conv_visits}."
    for i in
    outliers_cov_visits.values]
p3_1_2 = document.add_paragraph(style='List Bullet')
p3_1_2.paragraph_format.line_spacing = 1.5
p3_1_2.style.font.size = Pt(12)
p3_1_2.paragraph_format.left_indent = Inches(0.5)
p3_1_2.add_run(outliers_formatter(conv_visits_outliers_list, 'Конверсия визитов из лендинга'))

aborted_outliers_list = [
    f"«{i[0]}» ({i[RK_LABELS.index('perc_aborted')]} %) что в {great_or_less_range(i[RK_LABELS.index('perc_aborted')], cur_rk_data.perc_aborted.mean())} чем в среднем ({cur_rk_data.perc_aborted.mean().round(2)} %) при этом медиана для данного параметра = {median_perc_abort}."
    for i in outliers_perc_abort.replace(np.nan, 0).values]
p3_1_3 = document.add_paragraph(style='List Bullet')
p3_1_3.paragraph_format.line_spacing = 1.5
p3_1_3.style.font.size = Pt(12)
p3_1_3.paragraph_format.left_indent = Inches(0.5)
p3_1_3.add_run(outliers_formatter(aborted_outliers_list, 'Доля отказов (относительно визитов)'))

time_outliers_list = [
    f"«{i[0]}» ({time_to_str(i[RK_LABELS.index('time')])}) что в {great_or_less_range(time_to_seconds(i[RK_LABELS.index('time')]), cur_rk_data.time.apply(time_to_seconds).mean())} чем в среднем ({time_to_str(seconds_to_time(cur_rk_data.time.apply(time_to_seconds).mean()))}) при этом медиана для данного параметра = {time_to_str(median_time)}."
    for i in outliers_time.values]
p3_1_4 = document.add_paragraph(style='List Bullet')
p3_1_4.paragraph_format.line_spacing = 1.5
p3_1_4.style.font.size = Pt(12)
p3_1_4.paragraph_format.left_indent = Inches(0.5)
p3_1_4.add_run(outliers_formatter(time_outliers_list, 'Время на сайте'))

# ПУНКТ 4 - ГРУППЫ
p4 = document.add_paragraph(f'Группы:',
                              style='List Number')
p4.paragraph_format.line_spacing = 1.5
p4.style.font.size = Pt(12)



document.save('automatic_report_test.docx')
